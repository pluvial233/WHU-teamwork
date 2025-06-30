from flask import Flask, render_template_string, request, redirect, url_for, session # type: ignore
from flask_sqlalchemy import SQLAlchemy # type: ignore
import os
import time
import webbrowser
import threading
from datetime import datetime, timedelta

app = Flask(__name__)
app.secret_key = 'book_management_system_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///books.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# 数据库模型定义
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    role = db.Column(db.String(20), default='user')  # 'admin' or 'user'
    borrow_records = db.relationship('BorrowRecord', backref='user', lazy=True)

class Book(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    author = db.Column(db.String(100), nullable=False)
    category = db.Column(db.String(50))
    isbn = db.Column(db.String(20), unique=True)
    stock = db.Column(db.Integer, default=1)
    borrow_records = db.relationship('BorrowRecord', backref='book', lazy=True)

class BorrowRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    book_id = db.Column(db.Integer, db.ForeignKey('book.id'), nullable=False)
    borrow_date = db.Column(db.DateTime, default=datetime.utcnow)
    due_date = db.Column(db.DateTime, default=lambda: datetime.utcnow() + timedelta(days=14))
    return_date = db.Column(db.DateTime, nullable=True)
    fine = db.Column(db.Float, default=0.0)

class Feedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    date = db.Column(db.DateTime, default=datetime.utcnow)

# 创建数据库表
with app.app_context():
    db.create_all()
    # 创建默认用户（如果不存在）
    if not User.query.filter_by(username='admin').first():
        admin = User(username='admin', password='admin', role='admin')
        db.session.add(admin)
    if not User.query.filter_by(username='user').first():
        test_user = User(username='user', password='user', role='user')
        db.session.add(test_user)
    
    # 添加测试图书数据
    if not Book.query.first():
        books = [
            Book(title='Python编程：从入门到实践', author='埃里克·马瑟斯', category='编程', isbn='9787115428028', stock=5),
            Book(title='算法导论', author='托马斯·H·科曼', category='计算机科学', isbn='9787111407010', stock=3),
            Book(title='红楼梦', author='曹雪芹', category='文学', isbn='9787020002207', stock=4),
            Book(title='三体', author='刘慈欣', category='科幻', isbn='9787536692930', stock=6),
            Book(title='人类简史', author='尤瓦尔·赫拉利', category='历史', isbn='9787508647357', stock=2)
        ]
        db.session.add_all(books)
    
    db.session.commit()

# 路由定义
# HTML模板字符串
INDEX_HTML = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>图书管理系统</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }
        .header { text-align: center; margin-bottom: 40px; }
        .nav { display: flex; justify-content: center; gap: 20px; margin-bottom: 40px; }
        .btn { padding: 10px 20px; text-decoration: none; background-color: #4CAF50; color: white; border-radius: 5px; }
        .btn:hover { background-color: #45a049; }
        .feature-section { margin-top: 60px; }
        .feature-card { border: 1px solid #ddd; padding: 20px; margin: 10px; border-radius: 8px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>图书管理系统</h1>
        <p>便捷的图书借阅与管理平台</p>
    </div>
    
    <div class="nav">
        <a href="{{ url_for('login') }}" class="btn">用户登录</a>
        <a href="{{ url_for('login') }}" class="btn">管理员登录</a>
    </div>
    
    <div class="feature-section">
        <h2>系统功能</h2>
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px;">
            <div class="feature-card">
                <h3>图书检索</h3>
                <p>快速查找各类图书资源</p>
            </div>
            <div class="feature-card">
                <h3>借阅管理</h3>
                <p>便捷的图书借阅与归还流程</p>
            </div>
            <div class="feature-card">
                <h3>用户中心</h3>
                <p>查看借阅历史与个人信息</p>
            </div>
            <div class="feature-card">
                <h3>库存管理</h3>
                <p>管理员图书入库与盘点</p>
            </div>
        </div>
    </div>
</body>
</html>'''

@app.route('/')
def index():
    return render_template_string(INDEX_HTML)

REGISTER_HTML = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>注册 - 图书管理系统</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 400px; margin: 50px auto; padding: 20px; }
        .register-container { border: 1px solid #ddd; padding: 30px; border-radius: 8px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }
        h2 { text-align: center; margin-bottom: 20px; }
        .form-group { margin-bottom: 15px; }
        label { display: block; margin-bottom: 5px; }
        input { width: 100%; padding: 8px; box-sizing: border-box; }
        .btn { width: 100%; padding: 10px; background-color: #4CAF50; color: white; border: none; border-radius: 5px; cursor: pointer; }
        .btn:hover { background-color: #45a049; }
        .error { color: red; text-align: center; margin-top: 10px; }
        .success { color: green; text-align: center; margin-top: 10px; }
        .login-link { text-align: center; margin-top: 15px; }
    </style>
</head>
<body>
    <div class="register-container">
        <h2>用户注册</h2>
        <form method="POST" action="{{ url_for('register') }}">
            <div class="form-group">
                <label for="username">用户名</label>
                <input type="text" id="username" name="username" required>
            </div>
            <div class="form-group">
                <label for="password">密码</label>
                <input type="password" id="password" name="password" required>
            </div>
            <button type="submit" class="btn">注册</button>
            {% if error %}
            <div class="error">{{ error }}</div>
            {% endif %}
            {% if success %}
            <div class="success">{{ success }}</div>
            {% endif %}
            <div class="login-link">已有账号？<a href="{{ url_for('login') }}">前往登录</a></div>
        </form>
    </div>
</body>
</html>'''

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        # 检查用户名是否已存在
        if User.query.filter_by(username=username).first():
            return render_template_string(REGISTER_HTML, error='用户名已存在')
        
        # 创建新用户
        new_user = User(username=username, password=password, role='user')
        db.session.add(new_user)
        db.session.commit()
        
        return render_template_string(REGISTER_HTML, success='注册成功，请登录')
    return render_template_string(REGISTER_HTML)

LOGIN_HTML = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>登录 - 图书管理系统</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 400px; margin: 50px auto; padding: 20px; }
        .login-container { border: 1px solid #ddd; padding: 30px; border-radius: 8px; box-shadow: 0 0 10px rgba(0,0,0,0.1); }
        h2 { text-align: center; margin-bottom: 20px; }
        .form-group { margin-bottom: 15px; }
        label { display: block; margin-bottom: 5px; }
        input { width: 100%; padding: 8px; box-sizing: border-box; }
        .btn { width: 100%; padding: 10px; background-color: #4CAF50; color: white; border: none; border-radius: 5px; cursor: pointer; }
        .btn:hover { background-color: #45a049; }
        .error { color: red; text-align: center; margin-top: 10px; }
    </style>
</head>
<body>
    <div class="login-container">
        <h2>用户登录</h2>
        <form method="POST" action="{{ url_for('login') }}">
            <div class="form-group">
                <label for="username">用户名</label>
                <input type="text" id="username" name="username" required>
            </div>
            <div class="form-group">
                <label for="password">密码</label>
                <input type="password" id="password" name="password" required>
            </div>
            <button type="submit" class="btn">登录</button>
            {% if error %}
            <div class="error">{{ error }}</div>
            {% endif %}
            <div class="register-link">没有账号？<a href="{{ url_for('register') }}">立即注册</a></div>
        </form>
    </div>
</body>
</html>'''

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            session['user_id'] = user.id
            session['role'] = user.role
            return redirect(url_for('dashboard'))
        # 登录失败，显示错误信息
        return render_template_string(LOGIN_HTML, error='用户名或密码错误')
    return render_template_string(LOGIN_HTML)

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('role', None)
    return redirect(url_for('index'))

DASHBOARD_HTML = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>用户中心 - 图书管理系统</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 1200px; margin: 0 auto; padding: 20px; }
        .header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 40px; }
        .nav { display: flex; gap: 20px; }
        .btn { padding: 8px 15px; text-decoration: none; background-color: #4CAF50; color: white; border-radius: 5px; border: none; cursor: pointer; }
        .btn:hover { background-color: #45a049; }
        .section { margin: 30px 0; padding: 20px; border: 1px solid #ddd; border-radius: 8px; }
        .admin-features { display: none; }
        .admin .admin-features { display: block; }
    </style>
</head>
<body class="{{ 'admin' if session.role == 'admin' else 'user' }}">
    <div class="header">
        <h1>图书管理系统 - 用户中心</h1>
        <div class="nav">
            <a href="{{ url_for('logout') }}" class="btn">退出登录</a>
        </div>
    </div>
    
    <div class="section">
        <h2>欢迎使用图书管理系统</h2>
        <p>当前用户: {{ user.username }} (已登录)</p>
        <p>用户角色: {{ session.role }}</p>
    </div>
    
    <div class="section" id="search">
        <h3>图书检索</h3>
        <form method="POST" action="{{ url_for('search_books') }}">
            <input type="text" name="search_query" placeholder="输入书名或作者" required>
            <button type="submit" class="btn">搜索</button>
        </form>
        {% if search_results and search_results|length > 0 %}
        <div class="search-results">
            <h4>搜索结果 (关键词: {{ search_query }}):</h4>
            <ul>
                {% for book in search_results %}
                <li>
                    {{ book.title }} (作者: {{ book.author }}) - 库存: {{ book.stock }}
                    {% if book.stock > 0 %}
                    <a href="{{ url_for('borrow_book', book_id=book.id) }}" class="btn">借阅</a>
                    {% else %}
                    <span style="color: red;">无库存</span>
                    {% endif %}
                </li>
                {% endfor %}
            </ul>
        </div>
        {% else %}
            {% if search_query %}
            <div class="search-results">
                <h4>搜索结果 (关键词: {{ search_query }}):</h4>
                <p style="color: #666;">无对应结果</p>
            </div>
            {% endif %}
        {% endif %}
    </div>

    <div class="section" id="borrow">
        <h3>{% if session.role == 'admin' %}所有借阅记录{% else %}我的借阅{% endif %}</h3>
        <ul>
            {% if session.role == 'admin' %}
            {% for record in all_borrow_records %}
            <li>
                用户: {{ record.user.username }} - 图书: {{ record.book.title }} - 借阅日期: {{ record.borrow_date.strftime('%Y-%m-%d') }}
                到期日期: {{ record.due_date.strftime('%Y-%m-%d') }}
                {% if record.return_date %}
                <span style="color: green;">已归还</span>
                {% else %}
                <span style="color: orange;">未归还</span>
                {% endif %}
            </li>
            {% endfor %}
        {% else %}
            {% for record in borrow_records %}
            <li>
                {{ record.book.title }} - 借阅日期: {{ record.borrow_date.strftime('%Y-%m-%d') }}
                到期日期: {{ record.due_date.strftime('%Y-%m-%d') }}
                {% if record.return_date %}
                <span style="color: green;">已归还</span>
                {% else %}
                <span style="color: orange;">未归还</span>
                <a href="{{ url_for('return_book', record_id=record.id) }}" class="btn">归还</a>
                {% endif %}
            </li>
            {% else %}
            <li>暂无借阅记录</li>
            {% endfor %}
        {% endif %}
        </ul>
    </div>

    <div class="section">
        <h3>快速功能
        <div style="display: flex; gap: 15px; flex-wrap: wrap;">
            <a href="#search" class="btn">图书检索</a>
            <a href="#borrow" class="btn">我的借阅</a>
            {% if session.role == 'admin' %}
            <a href="#inventory" class="btn">库存管理</a>
            <a href="#users" class="btn">用户管理</a>
            {% endif %}
        </div>
    </div>
    
    <div class="section admin-features">
        <h3>管理员功能区</h3>
        <p>图书入库 | 编辑图书 | 删除图书 | 数据备份</p>
    </div>
</body>
</html>'''

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # 获取当前用户
    current_user = User.query.get(session['user_id'])
    
    # 根据角色获取不同数据
    if session['role'] == 'admin':
        # 管理员可以查看所有借阅记录
        all_borrow_records = BorrowRecord.query.all()
        return render_template_string(DASHBOARD_HTML, user=current_user, all_borrow_records=all_borrow_records, session=session)
    else:
        # 普通用户只能查看自己的借阅记录
        user_borrow_records = BorrowRecord.query.filter_by(user_id=session['user_id']).all()
        return render_template_string(DASHBOARD_HTML, user=current_user, borrow_records=user_borrow_records, session=session)

@app.route('/search_books', methods=['POST'])
def search_books():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    query = request.form['search_query']
    # 搜索书名或作者包含查询词的图书
    search_results = Book.query.filter(
        db.or_(Book.title.like(f'%{query}%'), Book.author.like(f'%{query}%'))
    ).all()
    # 获取当前用户
    current_user = User.query.get(session['user_id'])
    
    # 根据角色获取不同数据
    if session['role'] == 'admin':
        all_borrow_records = BorrowRecord.query.all()
        return render_template_string(DASHBOARD_HTML, user=current_user, search_results=search_results, all_borrow_records=all_borrow_records, search_query=query, session=session)
    else:
        user_borrow_records = BorrowRecord.query.filter_by(user_id=session['user_id']).all()
        return render_template_string(DASHBOARD_HTML, user=current_user, search_results=search_results, borrow_records=user_borrow_records, search_query=query, session=session)

@app.route('/borrow_book/<int:book_id>')
def borrow_book(book_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    book = Book.query.get_or_404(book_id)
    # 检查库存
    if book.stock <= 0:
        return redirect(url_for('dashboard'))
    # 创建借阅记录
    new_record = BorrowRecord(user_id=session['user_id'], book_id=book_id)
    db.session.add(new_record)
    # 减少库存
    book.stock -= 1
    db.session.commit()
    return redirect(url_for('dashboard'))

@app.route('/return_book/<int:record_id>')
def return_book(record_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    record = BorrowRecord.query.get_or_404(record_id)
    # 检查是否是当前用户的借阅记录
    if record.user_id != session['user_id']:
        return redirect(url_for('dashboard'))
    # 更新归还日期
    record.return_date = datetime.utcnow()
    # 增加库存
    book = Book.query.get(record.book_id)
    book.stock += 1
    db.session.commit()
    return redirect(url_for('dashboard'))

def open_browser_after_delay():
    # 等待服务器启动
    time.sleep(2)
    webbrowser.open('http://127.0.0.1:5000/')

def generate_system_docs():
    from docx import Document
    
    # 创建文档对象
    document = Document()
    
    # 设置标题
    document.add_heading('图书管理系统设计说明书', 0)
    
    # 1. 系统体系架构
    document.add_heading('1. 系统体系架构', level=1)
    paragraph = document.add_paragraph('本系统采用三层架构设计，基于Flask框架开发：')
    paragraph.add_run('\n- 表示层：').bold = True
    paragraph.add_run('使用HTML模板和CSS实现用户界面，包括登录、注册、图书检索和用户中心等页面')
    paragraph.add_run('\n- 业务逻辑层：').bold = True
    paragraph.add_run('通过Flask路由处理用户请求，实现图书借阅、归还、搜索等核心功能')
    paragraph.add_run('\n- 数据访问层：').bold = True
    paragraph.add_run('使用SQLAlchemy ORM框架与SQLite数据库交互，管理用户、图书和借阅记录数据')
    
    document.add_heading('系统技术栈', level=2)
    tech_table = document.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'
    tech_hdr_cells = tech_table.rows[0].cells
    tech_hdr_cells[0].text = '技术类别'
    tech_hdr_cells[1].text = '具体技术'
    tech_rows = [
        ('后端', 'Python 3.7+, Flask 2.0+, SQLAlchemy'),
        ('前端', 'HTML5, CSS3, Jinja2模板引擎'),
        ('数据库', 'SQLite')
    ]
    for tech_type, tech_details in tech_rows:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = tech_type
        row_cells[1].text = tech_details
    
    # 2. 系统功能结构
    document.add_heading('2. 系统功能结构', level=1)
    func_structure = document.add_paragraph('系统功能按用户角色分为：')
    
    func_structure.add_run('\n\n2.1 公共功能').bold = True
    func_structure.add_run('\n- 用户注册/登录')
    func_structure.add_run('\n- 图书检索（按书名、作者）')
    
    func_structure.add_run('\n\n2.2 普通用户功能').bold = True
    func_structure.add_run('\n- 查看个人借阅历史')
    func_structure.add_run('\n- 借阅图书')
    func_structure.add_run('\n- 归还图书')
    
    func_structure.add_run('\n\n2.3 管理员功能').bold = True
    func_structure.add_run('\n- 查看所有借阅记录')
    func_structure.add_run('\n- 图书库存管理')
    func_structure.add_run('\n- 用户管理')
    func_structure.add_run('\n- 数据备份')
    
    # 3. 系统用例时序图
    document.add_heading('3. 系统用例时序图', level=1)
    use_case = document.add_paragraph('以图书借阅流程为例：')
    use_case.add_run('\n1. 用户在登录状态下提交图书借阅请求')
    use_case.add_run('\n2. 系统验证用户身份和图书库存状态')
    use_case.add_run('\n3. 系统创建借阅记录并更新图书库存')
    use_case.add_run('\n4. 返回操作结果给用户')
    
    # 4. 复杂功能的算法设计
    document.add_heading('4. 复杂功能的算法设计', level=1)
    
    document.add_heading('4.1 图书搜索算法', level=2)
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = 'function search_books(keyword):'
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = '    results = []'
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = '    for book in all_books:'
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = '        if keyword in book.title or keyword in book.author:'
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = '            add book to results'
    code_paragraph = document.add_paragraph()
    code_paragraph.style = 'List Bullet'
    code_paragraph.text = '    return results'
    
    # 5. 面向对象方法类图设计
    document.add_heading('5. 面向对象方法类图设计', level=1)
    class_design = document.add_paragraph('核心类定义：')
    class_design.add_run('\n- User类：').bold = True
    class_design.add_run('id, username, password, role, borrow_records')
    class_design.add_run('\n- Book类：').bold = True
    class_design.add_run('id, title, author, category, isbn, stock, borrow_records')
    class_design.add_run('\n- BorrowRecord类：').bold = True
    class_design.add_run('id, user_id, book_id, borrow_date, due_date, return_date, fine')
    
    class_design.add_run('\n\n类关系：')
    class_design.add_run('\n- User与BorrowRecord是一对多关系')
    class_design.add_run('\n- Book与BorrowRecord是一对多关系')
    
    # 6. 接口设计
    document.add_heading('6. 接口设计', level=1)
    interface_table = document.add_table(rows=1, cols=2)
    interface_table.style = 'Table Grid'
    interface_hdr_cells = interface_table.rows[0].cells
    interface_hdr_cells[0].text = '请求方法'
    interface_hdr_cells[1].text = '接口路径及说明'
    interface_rows = [
        ('GET', '/login - 用户登录页面'),
        ('POST', '/login - 提交登录信息'),
        ('GET', '/register - 用户注册页面'),
        ('POST', '/register - 提交注册信息'),
        ('GET', '/dashboard - 用户中心页面'),
        ('POST', '/search_books - 图书搜索接口'),
        ('GET', '/borrow_book/<book_id> - 借阅图书接口'),
        ('GET', '/return_book/<record_id> - 归还图书接口')
    ]
    for method, path in interface_rows:
        row_cells = interface_table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = path
    
    # 7. 数据库物理设计
    document.add_heading('7. 数据库物理设计', level=1)
    document.add_paragraph('系统使用SQLite数据库，主要表结构：')
    
    # users表
    document.add_heading('users表', level=2)
    users_table = document.add_table(rows=1, cols=3)
    users_table.style = 'Table Grid'
    users_hdr_cells = users_table.rows[0].cells
    users_hdr_cells[0].text = '字段名'
    users_hdr_cells[1].text = '类型'
    users_hdr_cells[2].text = '约束'
    users_rows = [
        ('id', 'INTEGER', 'PRIMARY KEY'),
        ('username', 'VARCHAR(50)', 'UNIQUE, NOT NULL'),
        ('password', 'VARCHAR(100)', 'NOT NULL'),
        ('role', 'VARCHAR(20)', 'DEFAULT \'user\'')
    ]
    for field, type_, constraint in users_rows:
        row_cells = users_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = type_
        row_cells[2].text = constraint
    
    # books表
    document.add_heading('books表', level=2)
    books_table = document.add_table(rows=1, cols=3)
    books_table.style = 'Table Grid'
    books_hdr_cells = books_table.rows[0].cells
    books_hdr_cells[0].text = '字段名'
    books_hdr_cells[1].text = '类型'
    books_hdr_cells[2].text = '约束'
    books_rows = [
        ('id', 'INTEGER', 'PRIMARY KEY'),
        ('title', 'VARCHAR(200)', 'NOT NULL'),
        ('author', 'VARCHAR(100)', 'NOT NULL'),
        ('category', 'VARCHAR(50)', ''),
        ('isbn', 'VARCHAR(20)', 'UNIQUE'),
        ('stock', 'INTEGER', 'DEFAULT 1')
    ]
    for field, type_, constraint in books_rows:
        row_cells = books_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = type_
        row_cells[2].text = constraint
    
    # borrow_records表
    document.add_heading('borrow_records表', level=2)
    borrow_table = document.add_table(rows=1, cols=3)
    borrow_table.style = 'Table Grid'
    borrow_hdr_cells = borrow_table.rows[0].cells
    borrow_hdr_cells[0].text = '字段名'
    borrow_hdr_cells[1].text = '类型'
    borrow_hdr_cells[2].text = '约束'
    borrow_rows = [
        ('id', 'INTEGER', 'PRIMARY KEY'),
        ('user_id', 'INTEGER', 'FOREIGN KEY, NOT NULL'),
        ('book_id', 'INTEGER', 'FOREIGN KEY, NOT NULL'),
        ('borrow_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'),
        ('due_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP+14天'),
        ('return_date', 'DATETIME', 'NULLABLE'),
        ('fine', 'FLOAT', 'DEFAULT 0.0')
    ]
    for field, type_, constraint in borrow_rows:
        row_cells = borrow_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = type_
        row_cells[2].text = constraint
    
    # 8. UI设计
    document.add_heading('8. UI设计', level=1)
    ui_design = document.add_paragraph('系统界面采用简洁现代的设计风格，主要页面包括：')
    
    ui_design.add_run('\n\n8.1 首页').bold = True
    ui_design.add_run('\n- 系统标题和功能介绍')
    ui_design.add_run('\n- 用户登录和管理员登录入口')
    ui_design.add_run('\n- 系统功能模块展示')
    
    ui_design.add_run('\n\n8.2 登录/注册页面').bold = True
    ui_design.add_run('\n- 简洁的表单设计')
    ui_design.add_run('\n- 用户名/密码输入框')
    ui_design.add_run('\n- 错误提示和状态反馈')
    
    ui_design.add_run('\n\n8.3 用户中心').bold = True
    ui_design.add_run('\n- 个人信息展示')
    ui_design.add_run('\n- 图书搜索功能')
    ui_design.add_run('\n- 借阅历史记录')
    ui_design.add_run('\n- 角色相关功能按钮')
    
    ui_design.add_run('\n\n8.4 管理员功能区').bold = True
    ui_design.add_run('\n- 所有借阅记录查看')
    ui_design.add_run('\n- 库存管理入口')
    ui_design.add_run('\n- 用户管理入口')
    
    ui_design.add_run('\n\n界面采用绿色为主色调，象征知识与成长，整体布局清晰，操作流程简单直观。')
    
    # 保存文档
    document.save('系统设计说明书.docx')
    print('Word文档生成成功：系统设计说明书.docx')

# 添加生成文档的路由
@app.route('/generate_docs')
def trigger_generate_docs():
    if 'user_id' not in session or session.get('role') != 'admin':
        return '无权限访问此功能', 403
    generate_system_docs()
    return '系统设计说明书已生成，请在项目根目录查看'

if __name__ == '__main__':
    # 在单独线程中启动浏览器，避免阻塞服务器启动
    threading.Thread(target=open_browser_after_delay).start()
    app.run(debug=True)